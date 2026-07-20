#!/usr/bin/env python3
"""根据天气JSON数据生成Word报告，包含图表"""

import json
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# ── 加载天气数据 ──
with open('/tmp/hangzhou_weather.json', 'r') as f:
    data = json.load(f)

city = data['city']
now = data['now']
daily = data['daily']

# ── 准备图表数据 ──
dates = [d['fxDate'][5:] for d in daily]  # MM-DD
temps_high = [int(d['tempMax']) for d in daily]
temps_low = [int(d['tempMin']) for d in daily]
precips = [float(d['precip']) for d in daily]
weather_texts = [d['textDay'] for d in daily]
humidities = [int(d['humidity']) for d in daily]

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'Noto Sans CJK SC', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# ── 图表1: 温度趋势 + 降水柱状 ──
fig, ax1 = plt.subplots(figsize=(10, 5))

color_temp_high = '#E74C3C'
color_temp_low = '#3498DB'
color_precip = '#2ECC71'

# 温度折线
ax1.plot(dates, temps_high, 'o-', color=color_temp_high, linewidth=2, markersize=6, label='最高温度 (°C)')
ax1.plot(dates, temps_low, 'o-', color=color_temp_low, linewidth=2, markersize=6, label='最低温度 (°C)')
# 填充温差区域
ax1.fill_between(range(len(dates)), temps_low, temps_high, alpha=0.12, color=color_temp_high)

# 温度数值标注
for i, (h, l) in enumerate(zip(temps_high, temps_low)):
    ax1.annotate(f'{h}°', (dates[i], h), textcoords="offset points", xytext=(0, 8),
                 fontsize=8, color=color_temp_high, ha='center')
    ax1.annotate(f'{l}°', (dates[i], l), textcoords="offset points", xytext=(0, -14),
                 fontsize=8, color=color_temp_low, ha='center')

ax1.set_ylabel('温度 (°C)', fontsize=11, color='#333')
ax1.set_ylim(min(temps_low) - 5, max(temps_high) + 5)
ax1.tick_params(axis='y', labelcolor='#555')

# 降水柱状图（右轴）
ax2 = ax1.twinx()
bars = ax2.bar(dates, precips, alpha=0.35, color=color_precip, label='降水量 (mm)', width=0.5)
ax2.set_ylabel('降水量 (mm)', fontsize=11, color=color_precip)
ax2.tick_params(axis='y', labelcolor=color_precip)
ax2.set_ylim(0, max(precips) * 2.5 if max(precips) > 0 else 10)

# 降水数值标注
for bar, p in zip(bars, precips):
    if p > 0:
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                 f'{p}mm', ha='center', va='bottom', fontsize=8, color=color_precip)

# 图例
lines1, labels1 = ax1.get_legend_handles_labels()
lines2, labels2 = ax2.get_legend_handles_labels()
ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=9, framealpha=0.8)

ax1.set_title(f'{city} 未来7天温度趋势与降水量', fontsize=14, fontweight='bold', pad=15)
ax1.grid(axis='y', alpha=0.3, linestyle='--')
fig.tight_layout()
fig.savefig('/tmp/weather_trend.png', dpi=150, bbox_inches='tight')
plt.close()

# ── 图表2: 湿度 + 天气状况 ──
fig, ax = plt.subplots(figsize=(10, 4))

colors_map = {
    '晴': '#F39C12', '多云': '#95A5A6', '阴': '#7F8C8D',
    '小雨': '#5DADE2', '中雨': '#2980B9', '大雨': '#1A5276',
    '暴雨': '#6C3483', '雷阵雨': '#E67E22', '雪': '#AED6F1',
}
bar_colors = [colors_map.get(t, '#BDC3C7') for t in weather_texts]
bars = ax.bar(dates, humidities, color=bar_colors, width=0.55, edgecolor='white', linewidth=0.5)

# 天气标注
for bar, t, h in zip(bars, weather_texts, humidities):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
            f'{t}\n{h}%', ha='center', va='bottom', fontsize=9, color='#333')

ax.set_ylabel('相对湿度 (%)', fontsize=11)
ax.set_ylim(0, 110)
ax.set_title(f'{city} 未来7天湿度与天气状况', fontsize=14, fontweight='bold', pad=15)
ax.grid(axis='y', alpha=0.3, linestyle='--')
fig.tight_layout()
fig.savefig('/tmp/weather_humidity.png', dpi=150, bbox_inches='tight')
plt.close()

print("图表已生成: /tmp/weather_trend.png, /tmp/weather_humidity.png")

# ══════════════════════════════════════════
# 生成 Word 报告
# ══════════════════════════════════════════
doc = Document()

# ── 页边距 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 标题页 ──
title = doc.add_heading(f'{city} 未来一周天气预报报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # 空行

# ── 实时天气概况 ──
doc.add_heading('一、实时天气概况', level=1)

summary_data = [
    ('当前天气', now.get('text', '-')),
    ('当前温度', f"{now.get('temp', '-')}°C（体感 {now.get('feelsLike', '-')}°C）"),
    ('风力风向', f"{now.get('windDir', '-')} {now.get('windScale', '-')}级"),
    ('相对湿度', f"{now.get('humidity', '-')}%"),
    ('大气压强', f"{round(int(now.get('pressure', 0)) / 10, 1) if now.get('pressure') else '-'} kPa"),
    ('能见度', f"{now.get('vis', '-')} km"),
]

summary_table = doc.add_table(rows=len(summary_data), cols=2, style='Light Grid Accent 1')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(summary_data):
    summary_table.cell(i, 0).text = label
    summary_table.cell(i, 1).text = value
    for cell in [summary_table.cell(i, 0), summary_table.cell(i, 1)]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ── 预报概览 ──
doc.add_heading('二、预报概览', level=1)
overview = doc.add_paragraph()
overview.add_run(f'根据和风天气数据，{city}未来一周（{daily[0]["fxDate"]} 至 {daily[-1]["fxDate"]}）').font.size = Pt(11)

# 天气趋势总结
rain_days = sum(1 for d in daily if '雨' in d['textDay'])
sunny_days = sum(1 for d in daily if d['textDay'] == '晴')
max_precip_day = max(daily, key=lambda d: float(d['precip']))
max_temp = max(temps_high)
min_temp = min(temps_low)

trend_text = (
    f'天气趋势：前{rain_days}天有降水过程，其中{daily[0]["fxDate"]}降水量最大为{max_precip_day["precip"]}mm；'
    f'后{sunny_days}天转为晴好天气。'
    f'气温范围 {min_temp}°C ~ {max_temp}°C，体感舒适。'
)
overview.add_run(trend_text).font.size = Pt(11)

doc.add_paragraph()

# ── 7天详细预报表格 ──
doc.add_heading('三、7天详细预报', level=1)

headers = ['日期', '天气', '最高温', '最低温', '降水量', '湿度', '风力', '紫外线']
table = doc.add_table(rows=len(daily) + 1, cols=len(headers), style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
for j, header in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# 数据行
for i, d in enumerate(daily):
    row_data = [
        f'{d["fxDate"][5:]}',
        f'{d["textDay"]}/{d["textNight"]}',
        f'{d["tempMax"]}°C',
        f'{d["tempMin"]}°C',
        f'{d["precip"]}mm',
        f'{d["humidity"]}%',
        f'{d["windDirDay"]} {d["windScaleDay"]}级',
        str(d.get('uvIndex', '-')),
    ]
    for j, val in enumerate(row_data):
        cell = table.cell(i + 1, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ── 图表1: 温度趋势 ──
doc.add_heading('四、温度趋势与降水量', level=1)
doc.add_picture('/tmp/weather_trend.png', width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

caption1 = doc.add_paragraph()
caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption1.add_run(f'图1: {city}未来7天温度趋势与降水量分布')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ── 图表2: 湿度 ──
doc.add_heading('五、湿度与天气状况', level=1)
doc.add_picture('/tmp/weather_humidity.png', width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

caption2 = doc.add_paragraph()
caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption2.add_run(f'图2: {city}未来7天湿度与天气状况')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ── 出行建议 ──
doc.add_heading('六、出行建议', level=1)

advices = []
if rain_days > 0:
    advices.append(f'未来{daily[0]["fxDate"]}-{daily[rain_days-1]["fxDate"]}有降水，建议随身携带雨具。')
if max_precip_day['precip'] and float(max_precip_day['precip']) > 10:
    advices.append(f'{float(max_precip_day["precip"])}mm/{max_precip_day["fxDate"][5:]}日降水量较大，注意防范城市内涝。')
if sunny_days > 2:
    advices.append(f'{daily[-sunny_days]["fxDate"]}起转为晴好天气，适合户外活动。')
max_uv = max(int(d.get('uvIndex', 0)) for d in daily)
if max_uv >= 8:
    advices.append(f'后几天紫外线指数高达{max_uv}，外出需做好防晒措施。')
if max_temp >= 30:
    advices.append('午后气温较高，注意防暑降温。')

for advice in advices:
    p = doc.add_paragraph(advice, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph()

# ── 数据来源 ──
doc.add_heading('七、数据来源', level=2)
source = doc.add_paragraph()
run = source.add_run('本报告数据来源于和风天气（QWeather）API，实时天气与7天预报数据。')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── 保存 ──
output_path = '/home/lizhan/project/long/tests/杭州未来一周天气预报报告.docx'
doc.save(output_path)
print(f'报告已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')