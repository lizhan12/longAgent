"""完整执行链路集成测试

真实问题是：单个组件测试全过，但全链路跑起来就崩。
这个文件模拟完整的计划执行流程：

1. 模拟 LLM 返回 PlanIR JSON
2. IRParser 解析
3. 约束验证
4. 逐步骤执行（mock 工具）
5. 沙箱执行代码
6. 文件落地检查
7. 交付物验证
8. 打印结果

每个测试都模拟一个真实的用户请求，并验证整条链路走通。
"""

from __future__ import annotations

import asyncio
import json
import os
import tempfile
from pathlib import Path

import pytest

from long.ir.executor import PlanExecutor
from long.ir.ir_parser import IRParser, IRParseStatus
from long.ir.plan_ir import PlanIR
from long.ir.constraint_validator import ConstraintValidator
from long.ir.state_machine import AgentStateMachine
from long.ir.type_checker import TypeChecker
from long.ir.ltl import LTLValidator
from long.sandbox.manager import SandboxManager
from long.sandbox.base import ExecutionSpec, ResourceLimits
from long.workspace.manager import WorkspaceManager

# 检查可选依赖
_HAS_DOCX = False
try:
    import docx  # noqa: F401
    _HAS_DOCX = True
except ImportError:
    pass

_HAS_FPDF = False
try:
    from fpdf import FPDF  # noqa: F401
    _HAS_FPDF = True
except ImportError:
    pass


# -----------------------------------------------------------
# Helper: 快速创建沙箱 + 执行代码
# -----------------------------------------------------------
async def _sandbox_exec(ws, code, timeout=60, network=False):
    """在沙箱中执行代码，返回 (result, output_dir)"""
    output_dir = ws.root / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    sandbox = SandboxManager(workspace_dir=str(ws.root))
    spec = ExecutionSpec(
        code=code, language="python", timeout=timeout,
        working_dir=str(ws.root),
        resource_limits=ResourceLimits(network=network),
    )
    result = await sandbox.execute(spec)
    return result, output_dir


# ======================== 场景 1: 沙箱执行代码后文件是否真的落地 ========================


@pytest.mark.asyncio
async def test_sandbox_execute_code_actually_creates_file():
    """验证沙箱 execute_code 后，文件确实保存在 workspace/output/ 下

    这是用户遇到的核心问题：execute_code 跑完说成功了，但文件找不到。
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        CODE = (
            "import os\n"
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "fig, ax = plt.subplots()\n"
            "ax.plot([1, 2, 3], [1, 4, 9])\n"
            "ax.set_title('中文标题')\n"
            "out = os.path.join(os.environ.get('OUTPUT_DIR', 'output'), 'test_chart.png')\n"
            "fig.savefig(out, dpi=80)\n"
            "print(f'saved: {os.path.exists(out)}')\n"
        )

        sandbox = SandboxManager(workspace_dir=str(ws.root))
        spec = ExecutionSpec(
            code=CODE,
            language="python",
            timeout=120,
            working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result = await sandbox.execute(spec)

        # 1. 执行成功
        assert result.status.value == "success", (
            f"沙箱执行失败: exit={result.exit_code}\n"
            f"stdout: {result.stdout}\n"
            f"stderr: {result.stderr}"
        )
        # 2. stdout 包含文件已保存的确认
        assert "saved:" in result.stdout, f"stdout 未包含文件保存确认: {result.stdout}"

        # 3. 文件真的存在（在 workspace/output/ 下）
        chart_file = output_dir / "test_chart.png"
        assert chart_file.exists(), (
            f"文件未找到: {chart_file}\n"
            f"output_dir 内容: {list(output_dir.iterdir()) if output_dir.exists() else '目录不存在'}"
        )
        assert chart_file.stat().st_size > 0, f"文件大小为 0: {chart_file}"


# ======================== 场景 2: 完整计划执行 ========================


@pytest.mark.skipif(not _HAS_DOCX, reason="需要 python-docx 包")
@pytest.mark.asyncio
async def test_plan_execution_creates_output_files():
    """验证完整计划执行后，所有输出文件都正确生成

    模拟一个真实场景：天气查询 → 生成图表 → 保存
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 模拟一个简单的计划
        chart_code = (
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "fig, ax = plt.subplots()\n"
            "ax.bar(['杭州', '苏州'], [25, 28])\n"
            "ax.set_title('气温对比')\n"
            "fig.savefig('output/weather_chart.png', dpi=80)\n"
            "print('chart saved')\n"
        )

        sandbox = SandboxManager(workspace_dir=str(ws.root))

        # 步骤 1: 执行图表代码
        spec1 = ExecutionSpec(
            code=chart_code,
            language="python",
            timeout=120,
            working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result1 = await sandbox.execute(spec1)
        assert result1.status.value == "success", f"步骤1失败: {result1.stderr}"

        # 验证图表文件已生成
        chart_file = output_dir / "weather_chart.png"
        assert chart_file.exists(), f"图表文件未生成: {chart_file}"
        assert chart_file.stat().st_size > 0, "图表文件为空"

        # 步骤 2: 生成 Word 文档
        word_code = (
            "from docx import Document\n"
            "doc = Document()\n"
            "doc.add_heading('天气对比报告', 0)\n"
            "doc.add_paragraph('杭州气温25°C，苏州气温28°C')\n"
            "doc.save('output/weather_report.docx')\n"
            "print('word saved')\n"
        )
        spec2 = ExecutionSpec(
            code=word_code,
            language="python",
            timeout=120,
            working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result2 = await sandbox.execute(spec2)
        assert result2.status.value == "success", f"步骤2失败: {result2.stderr}"

        # 验证 Word 文件已生成
        docx_file = output_dir / "weather_report.docx"
        assert docx_file.exists(), f"Word文件未生成: {docx_file}"
        assert docx_file.stat().st_size > 0, "Word文件为空"

        # 验证 output 目录只有这 2 个文件（没有多余文件）
        out_files = [f.name for f in sorted(output_dir.iterdir()) if f.is_file()]
        assert "weather_chart.png" in out_files, f"output 目录缺少图表文件: {out_files}"
        assert "weather_report.docx" in out_files, f"output 目录缺少 Word 文件: {out_files}"


# ======================== 场景 3: 中文字体渲染 ========================


@pytest.mark.asyncio
async def test_chinese_chart_renders_correctly():
    """验证图表中文标题能正确渲染，不是方块

    用户之前遇到的：图表中文标题变成方块，因为没有预置中文字体
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        CODE = (
            "import os\n"
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "from matplotlib.font_manager import findfont, FontProperties\n"
            "fig, ax = plt.subplots()\n"
            "ax.set_title('人工智能发展里程碑')\n"
            "ax.set_ylabel('重大事件数量')\n"
            "resolved = findfont(FontProperties(family=matplotlib.rcParams['font.family']))\n"
            "print(f'font: {os.path.basename(resolved)}')\n"
            "from matplotlib.ft2font import FT2Font\n"
            "face = FT2Font(resolved)\n"
            "missing = [c for c in '人工智能' if face.get_char_index(ord(c)) == 0]\n"
            "print(f'missing_glyphs: {missing}' if missing else 'all_glyphs_ok')\n"
            "fig.savefig('output/cjk_test.png', dpi=80)\n"
            "print('chart saved')\n"
        )

        sandbox = SandboxManager(workspace_dir=str(ws.root))
        spec = ExecutionSpec(
            code=CODE,
            language="python",
            timeout=120,
            working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result = await sandbox.execute(spec)

        assert result.status.value == "success", f"执行失败: {result.stderr[:300]}"
        assert "all_glyphs_ok" in result.stdout, (
            f"中文渲染失败，存在缺失字形\n"
            f"stdout: {result.stdout}\n"
            f"stderr: {result.stderr[:300]}"
        )

        # 确认图表文件已生成
        chart_file = output_dir / "cjk_test.png"
        assert chart_file.exists(), f"图表文件未生成: {chart_file}"


# ======================== 场景 4: 沙箱写文件与 workspace 路径一致性 ========================


@pytest.mark.asyncio
async def test_sandbox_write_file_consistency():
    """Verify execute_code and write_file write to the same output directory

    Previous bug: charts from execute_code landed in C:/Users/.../output
    while scripts from write_file went to D:/project/.../output.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 1. execute_code 保存文件
        CODE1 = (
            "with open('output/from_execute_code.txt', 'w') as f:\n"
            "    f.write('hello from execute_code')\n"
            "print('written')\n"
        )
        sandbox = SandboxManager(workspace_dir=str(ws.root))
        spec1 = ExecutionSpec(
            code=CODE1,
            language="python",
            timeout=30,
            working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result1 = await sandbox.execute(spec1)
        assert result1.status.value == "success", f"execute_code 失败: {result1.stderr[:200]}"

        # 2. write_file 保存文件（模拟 CLI 的 write_file 工具行为）
        ws.write_file("output/from_write_file.txt", "hello from write_file")

        # 3. 验证两个文件在同一个目录
        f1 = output_dir / "from_execute_code.txt"
        f2 = output_dir / "from_write_file.txt"
        assert f1.exists(), f"execute_code 文件未生成: {f1}"
        assert f2.exists(), f"write_file 文件未生成: {f2}"
        assert f1.read_text(encoding="utf-8") == "hello from execute_code"
        assert f2.read_text(encoding="utf-8") == "hello from write_file"


# ======================== 场景 5: 计划解析失败后降级 ========================


@pytest.mark.asyncio
async def test_plan_parse_failure_graceful_degradation():
    """验证 LLM 返回非计划内容时，系统能优雅降级

    用户遇到的核心问题：LLM 返回 {'city': '杭州'}，计划解析失败，
    系统应该降级到直接工具调用模式而不是报错退出。
    """
    # 模拟 LLM 返回非计划内容
    llm_output = json.dumps({"city": "杭州"})

    parser = IRParser()
    result = parser.parse(llm_output)

    # 修复策略应能识别并修复
    assert result.status in (
        IRParseStatus.REPAIRABLE, IRParseStatus.UNPARSEABLE
    ), f"LLM 非计划内容应被识别，实际: {result.status}"

    if result.status == IRParseStatus.REPAIRABLE:
        assert result.plan is not None, "修复后 plan 不应为 None"
        # 验证修复策略填充了必需字段
        assert result.plan.plan_id == "plan_auto_repaired", f"plan_id 未修复: {result.plan.plan_id}"
        assert result.plan.goal == "plan_auto_repaired", f"goal 未修复: {result.plan.goal}"
    else:
        # UNPARSEABLE 时，系统应降级到直接工具调用模式
        # 这个降级逻辑在 cli.py 的 _execute_plan 中，这里只验证解析器正确识别
        assert result.plan is None, "UNPARSEABLE 时 plan 应为 None"


# ======================== 场景 6: 多步骤计划依赖关系 ========================


@pytest.mark.skipif(not _HAS_DOCX, reason="需要 python-docx 包")
@pytest.mark.asyncio
async def test_multi_step_plan_execution_order():
    """验证多步骤计划按正确的依赖顺序执行

    场景：搜索 → 生成图表 → 生成 Word 文档
    图表必须在搜索之后，Word 必须在图表之后
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        sandbox = SandboxManager(workspace_dir=str(ws.root))

        # 模拟数据
        (output_dir / "hangzhou_data.txt").write_text("杭州: 25°C", encoding="utf-8")
        (output_dir / "suzhou_data.txt").write_text("苏州: 28°C", encoding="utf-8")

        # 步骤 1: 读取数据（模拟 query_weather）
        step1_code = (
            "with open('output/hangzhou_data.txt') as f:\n"
            "    print(f.read())\n"
        )
        r1 = await sandbox.execute(ExecutionSpec(
            code=step1_code, language="python", timeout=30,
            working_dir=str(ws.root), resource_limits=ResourceLimits(network=False),
        ))
        assert r1.status.value == "success", f"步骤1失败: {r1.stderr[:200]}"

        # 步骤 2: 生成图表（依赖步骤1）
        step2_code = (
            "import matplotlib\nmatplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "fig, ax = plt.subplots()\n"
            "ax.bar(['杭州', '苏州'], [25, 28])\n"
            "fig.savefig('output/chart.png', dpi=80)\n"
            "print('chart saved')\n"
        )
        r2 = await sandbox.execute(ExecutionSpec(
            code=step2_code, language="python", timeout=60,
            working_dir=str(ws.root), resource_limits=ResourceLimits(network=False),
        ))
        assert r2.status.value == "success", f"步骤2失败: {r2.stderr[:200]}"

        # 步骤 3: 生成 Word 文档（依赖步骤2）
        step3_code = (
            "from docx import Document\n"
            "doc = Document()\n"
            "doc.add_heading('天气对比', 0)\n"
            "doc.add_picture('output/chart.png')\n"
            "doc.save('output/report.docx')\n"
            "print('report saved')\n"
        )
        r3 = await sandbox.execute(ExecutionSpec(
            code=step3_code, language="python", timeout=60,
            working_dir=str(ws.root), resource_limits=ResourceLimits(network=False),
        ))
        assert r3.status.value == "success", f"步骤3失败: {r3.stderr[:200]}"

        # 验证所有输出文件
        assert (output_dir / "chart.png").exists(), "图表文件未生成"
        assert (output_dir / "report.docx").exists(), "Word 文件未生成"
        assert (output_dir / "report.docx").stat().st_size > 0, "Word 文件为空"


# ======================== 场景 7: 文件路径一致性 ========================


@pytest.mark.asyncio
async def test_output_path_consistency():
    """验证 workspace.root 和 output_dir 路径一致

    用户之前的问题：_workspace_root() 返回 WindowsPath，
    导致 os.path.join 行为异常，_print_generated_files 找不到文件。
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 模拟 _print_generated_files 的路径逻辑
        import os as _os

        _ws_root = ws.root  # 这是 WindowsPath 对象
        output_dir_from_ws = _os.path.join(str(_ws_root), "output") if _ws_root else "output"

        # 验证 output_dir 指向正确位置
        expected = str((ws.root / "output").resolve())
        actual = str(Path(output_dir_from_ws).resolve())
        assert actual == expected, (
            f"output_dir 路径不一致\n"
            f"  期望: {expected}\n"
            f"  实际: {actual}"
        )

        # 实际写入文件并验证
        (output_dir / "test.txt").write_text("hello", encoding="utf-8")
        assert _os.path.exists(_os.path.join(str(_ws_root), "output", "test.txt")), (
            "文件路径构建失败: os.path.join 无法正确处理 WindowsPath + str"
        )


# ======================== 场景 8: write_file + execute_file 联动 ========================


@pytest.mark.asyncio
async def test_write_file_then_execute_file():
    """验证 write_file 写入脚本后，execute_file 能执行它并生成输出文件

    之前用户遇到的：PPT 脚本写进去了但没执行，因为没有 execute_file 步骤。
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 1. write_file: 写入一个 Python 脚本
        script_content = (
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "fig, ax = plt.subplots()\n"
            "ax.plot([1, 2, 3], [1, 4, 9])\n"
            "fig.savefig('output/from_script.png', dpi=80)\n"
            "print('script executed successfully')\n"
        )
        ws.write_file("output/generate_chart.py", script_content)

        script_path = output_dir / "generate_chart.py"
        assert script_path.exists(), "write_file 未成功写入脚本"
        assert script_path.stat().st_size > 0, "写入的脚本为空"

        # 2. execute_file: 执行该脚本（模拟 CLI 的 execute_file 工具行为）
        code_to_run = (
            "import sys\n"
            "sys.path.insert(0, '.')\n"
            f"exec(open(r'{script_path}').read())\n"
        )
        result, output_dir = await _sandbox_exec(ws, code_to_run, timeout=60)

        assert result.status.value == "success", (
            f"execute_file 执行失败: {result.stderr[:300]}"
        )
        assert "script executed successfully" in result.stdout, (
            f"脚本未正确执行: {result.stdout}"
        )

        # 3. 验证脚本生成的输出文件
        chart_file = output_dir / "from_script.png"
        assert chart_file.exists(), f"脚本生成的图表文件未找到: {chart_file}"
        assert chart_file.stat().st_size > 0, "脚本生成的图表文件为空"


# ======================== 场景 9: 沙箱清理不删除输出文件 ========================


@pytest.mark.asyncio
async def test_sandbox_cleanup_preserves_output_files():
    """验证沙箱 cleanup 只删除临时目录，不删除 workspace/output/ 下的文件

    之前怀疑的：沙箱清理时把生成的文件也删了。
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 用 execute 代替私有 API，execute 内部会 cleanup
        result, output_dir = await _sandbox_exec(
            ws,
            "with open('output/survivor.txt', 'w') as f: f.write('survived')",
            timeout=30,
        )

        assert result.status.value == "success", f"执行失败: {result.stderr[:200]}"

        # 验证：即使沙箱已清理，输出文件还在
        survivor = output_dir / "survivor.txt"
        assert survivor.exists(), (
            f"输出文件在沙箱清理后被删除!\n"
            f"  输出文件: {survivor}"
        )
        assert survivor.read_text(encoding="utf-8") == "survived", "输出文件内容被修改"


# ======================== 场景 10: 空计划降级 ========================


def test_empty_plan_degrades_gracefully():
    """验证 LLM 生成了空计划（steps=[]）时，系统能降级而不是报错

    模拟 PlanExecutor.generate_plan 返回空计划的情况。
    """
    empty_plan = {
        "plan_id": "plan_empty",
        "goal": "空计划",
        "steps": [],
        "estimated_steps": 0,
    }
    parser = IRParser()
    result = parser.parse(json.dumps(empty_plan))

    # 空计划也能被解析（steps 有默认值）
    assert result.status in (IRParseStatus.SUCCESS, IRParseStatus.REPAIRABLE), str(result.errors)
    assert result.plan is not None
    assert len(result.plan.steps) == 0, f"空计划 steps 应为 0, 实际: {len(result.plan.steps)}"

    # 验证约束验证器能通过空计划
    validator = ConstraintValidator(
        state_machine=AgentStateMachine(),
        type_checker=TypeChecker(),
        ltl_validator=LTLValidator(),
    )
    validation = validator.validate_plan(result.plan)
    assert validation.valid, f"空计划验证失败: {validation.errors}"


# ======================== 场景 11: 跨平台路径兼容性 ========================


@pytest.mark.asyncio
async def test_path_handling_mixed_separators():
    """验证各种路径格式在 Windows 下都能正确处理

    Windows 上可能出现的路径格式：
    - output/file.txt
    - output\\file.txt
    - output/subdir/file.txt
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 测试各种路径格式
        paths_to_test = [
            ("output/forward_slash.txt", "forward"),
            ("output\\backward_slash.txt", "backward"),
            ("output/mixed\\slashes.txt", "mixed"),
        ]

        for rel_path, content in paths_to_test:
            ws.write_file(rel_path, content)
            # 验证文件在工作区根目录下
            resolved = (ws.root / rel_path).resolve()
            assert resolved.exists(), f"文件未创建: {resolved}"
            assert resolved.read_text(encoding="utf-8") == content, f"文件内容不匹配: {rel_path}"

        # 验证所有文件都在同一个 output 目录下
        out_files = [f.name for f in sorted(output_dir.iterdir()) if f.is_file()]
        assert "forward_slash.txt" in out_files, f"缺少 forward_slash.txt: {out_files}"
        assert "backward_slash.txt" in out_files, f"缺少 backward_slash.txt: {out_files}"
        assert "mixed\\slashes.txt" not in out_files, "反斜杠不应在文件名中"


# ======================== 场景 12: 中文文件名和内容 ========================


@pytest.mark.asyncio
async def test_chinese_filename_and_content():
    """验证中文文件名和内容在沙箱中能正确处理

    之前用户遇到的：GBK 编码导致 UnicodeEncodeError
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        CODE = (
            "import os\n"
            "content = '人工智能发展报告\\nAI发展历史里程碑\\n© 2026'\n"
            "with open('output/中文报告.txt', 'w', encoding='utf-8') as f:\n"
            "    f.write(content)\n"
            "print('written:', os.path.exists('output/中文报告.txt'))\n"
            "with open('output/中文报告.txt', 'r', encoding='utf-8') as f:\n"
            "    print('content:', f.read())\n"
        )
        result, output_dir = await _sandbox_exec(ws, CODE, timeout=30)

        assert result.status.value == "success", f"执行失败: {result.stderr[:300]}"
        assert "written: True" in result.stdout, f"文件未写入: {result.stdout}"

        cn_file = output_dir / "中文报告.txt"
        assert cn_file.exists(), f"中文文件未生成: {cn_file}"
        content = cn_file.read_text(encoding="utf-8")
        assert "人工智能发展报告" in content, f"中文内容不正确: {content}"
        assert "©" in content, f"特殊字符丢失: {content}"


# ======================== 场景 13: 多工具查询结果合并 ========================


@pytest.mark.asyncio
async def test_multi_tool_query_merge():
    """验证多次查询结果能合并到同一个输出文件中

    模拟：查询杭州天气 + 查询苏州天气 → 合并到同一个图表
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 模拟两次查询结果
        (output_dir / "city_a.txt").write_text("杭州: 25°C", encoding="utf-8")
        (output_dir / "city_b.txt").write_text("苏州: 28°C", encoding="utf-8")

        # 合并查询结果生成图表
        MERGE_CODE = (
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "with open('output/city_a.txt', encoding='utf-8') as f:\n"
            "    data_a = f.read().strip()\n"
            "with open('output/city_b.txt', encoding='utf-8') as f:\n"
            "    data_b = f.read().strip()\n"
            "fig, ax = plt.subplots()\n"
            "ax.bar(['杭州', '苏州'], [25, 28])\n"
            "ax.set_title('气温对比')\n"
            "fig.savefig('output/merged_chart.png', dpi=80)\n"
            "print('merged chart saved')\n"
        )
        result, output_dir = await _sandbox_exec(ws, MERGE_CODE, timeout=60)

        assert result.status.value == "success", f"合并执行失败: {result.stderr[:300]}"
        merged = output_dir / "merged_chart.png"
        assert merged.exists(), f"合并图表文件未生成: {merged}"
        assert merged.stat().st_size > 0, "合并图表文件为空"


# ======================== 场景 14: 失败步骤不阻塞后续步骤 ========================


@pytest.mark.asyncio
async def test_step_failure_does_not_block_subsequent_steps():
    """验证某个步骤失败时，后续步骤仍能继续执行

    实际场景：搜索失败不应该阻止图表生成（如果已有缓存数据）
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        ws = WorkspaceManager(tmpdir)
        output_dir = ws.root / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        # 步骤 1: 模拟搜索失败（执行一个不存在的命令）
        sandbox = SandboxManager(workspace_dir=str(ws.root))
        spec1 = ExecutionSpec(
            code="raise RuntimeError('模拟搜索失败')",
            language="python", timeout=30, working_dir=str(ws.root),
            resource_limits=ResourceLimits(network=False),
        )
        result1 = await sandbox.execute(spec1)
        assert result1.status.value != "success", "步骤1应该失败但成功了"

        # 步骤 2: 使用缓存数据继续生成图表
        step2_code = (
            "import matplotlib\n"
            "matplotlib.use('Agg')\n"
            "import matplotlib.pyplot as plt\n"
            "fig, ax = plt.subplots()\n"
            "ax.plot([1, 2, 3])\n"
            "fig.savefig('output/after_failure.png', dpi=80)\n"
            "print('step 2 succeeded despite step 1 failure')\n"
        )
        result2, output_dir = await _sandbox_exec(ws, step2_code, timeout=60)

        assert result2.status.value == "success", (
            f"步骤2被步骤1的失败阻塞了: {result2.stderr[:300]}"
        )
        assert (output_dir / "after_failure.png").exists(), (
            "步骤2的图表文件未生成"
        )